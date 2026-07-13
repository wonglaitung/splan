#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Converter
将 Markdown 文件转换为 Word 文档，保留完整格式
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 缺少 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)


class MarkdownToWordConverter:
    """Markdown 转 Word 转换器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def convert(self, md_file_path, output_path=None):
        """转换 Markdown 文件为 Word 文档"""
        md_path = Path(md_file_path)
        
        if not md_path.exists():
            raise FileNotFoundError(f"文件不存在: {md_file_path}")
        
        # 读取 Markdown 内容
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 解析并转换
        self.parse_markdown(content)
        
        # 确定输出路径
        if output_path is None:
            output_path = md_path.with_suffix('.docx')
        
        # 保存文档
        self.doc.save(output_path)
        print(f"✓ 转换成功: {output_path}")
        return output_path
    
    def parse_markdown(self, content):
        """解析 Markdown 内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 标题
            if line.startswith('#'):
                self.add_heading(line)
                i += 1
            # 代码块
            elif line.startswith('```'):
                i = self.add_code_block(lines, i)
            # 表格
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                i = self.add_table(lines, i)
            # 水平线
            elif re.match(r'^[-*_]{3,}\s*$', line):
                self.add_horizontal_line()
                i += 1
            # 无序列表
            elif line.strip().startswith(('-', '*', '+')):
                self.add_bullet_list(lines, i)
                i += 1
            # 有序列表
            elif re.match(r'^\d+\.', line.strip()):
                self.add_numbered_list(lines, i)
                i += 1
            # 引用
            elif line.startswith('>'):
                self.add_quote(line)
                i += 1
            # 空行
            elif not line.strip():
                i += 1
            # 普通段落
            else:
                self.add_paragraph(line)
                i += 1
    
    def add_heading(self, line):
        """添加标题"""
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            heading = self.doc.add_heading(text, level=level)
            # 设置中文字体
            for run in heading.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def add_paragraph(self, line):
        """添加段落，处理行内格式"""
        para = self.doc.add_paragraph()
        self.add_inline_formatting(para, line)
    
    def add_inline_formatting(self, para, text):
        """处理行内格式（粗体、斜体、代码）"""
        # 简单的行内格式处理
        parts = []
        current = ''
        i = 0
        
        while i < len(text):
            # 粗体
            if text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ''
                j = i + 2
                while j < len(text) and text[j:j+2] != '**':
                    j += 1
                if j < len(text):
                    parts.append(('bold', text[i+2:j]))
                    i = j + 2
                    continue
            # 斜体
            elif text[i] == '*' and (i == 0 or text[i-1] != '*'):
                if current:
                    parts.append(('normal', current))
                    current = ''
                j = i + 1
                while j < len(text) and text[j] != '*':
                    j += 1
                if j < len(text):
                    parts.append(('italic', text[i+1:j]))
                    i = j + 1
                    continue
            # 行内代码
            elif text[i] == '`':
                if current:
                    parts.append(('normal', current))
                    current = ''
                j = i + 1
                while j < len(text) and text[j] != '`':
                    j += 1
                if j < len(text):
                    parts.append(('code', text[i+1:j]))
                    i = j + 1
                    continue
            
            current += text[i]
            i += 1
        
        if current:
            parts.append(('normal', current))
        
        # 添加到段落
        for fmt, txt in parts:
            run = para.add_run(txt)
            if fmt == 'bold':
                run.bold = True
            elif fmt == 'italic':
                run.italic = True
            elif fmt == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
    
    def add_code_block(self, lines, start_idx):
        """添加代码块"""
        # 找到代码块结束位置
        i = start_idx + 1
        code_lines = []
        
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        # 添加代码块
        if code_lines:
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            
            code_text = '\n'.join(code_lines)
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        return i + 1
    
    def add_table(self, lines, start_idx):
        """添加表格"""
        # 收集表格行
        table_lines = []
        i = start_idx
        
        while i < len(lines) and '|' in lines[i]:
            table_lines.append(lines[i])
            i += 1
        
        if len(table_lines) < 2:
            return i
        
        # 解析表格
        rows = []
        for line in table_lines:
            # 跳过分隔行
            if re.match(r'^[\|\-\s:]+$', line):
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
        
        if not rows:
            return i
        
        # 创建表格
        num_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # 填充表格
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    # 表头加粗
                    if row_idx == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
        
        return i
    
    def add_horizontal_line(self):
        """添加水平线"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run('_' * 80)
        run.font.color.rgb = RGBColor(200, 200, 200)
    
    def add_bullet_list(self, lines, idx):
        """添加无序列表"""
        line = lines[idx]
        text = re.sub(r'^[\s\-\*\+]+', '', line).strip()
        if text:
            para = self.doc.add_paragraph(text, style='List Bullet')
    
    def add_numbered_list(self, lines, idx):
        """添加有序列表"""
        line = lines[idx]
        text = re.sub(r'^\d+\.\s*', '', line).strip()
        if text:
            para = self.doc.add_paragraph(text, style='List Number')
    
    def add_quote(self, line):
        """添加引用"""
        text = line.lstrip('>').strip()
        if text:
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python md_to_word.py <input.md> [--output <output.docx>]")
        print("示例: python md_to_word.py README.md")
        print("      python md_to_word.py README.md --output report.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = None
    
    # 解析参数
    if '--output' in sys.argv or '-o' in sys.argv:
        try:
            output_idx = sys.argv.index('--output') if '--output' in sys.argv else sys.argv.index('-o')
            output_file = sys.argv[output_idx + 1]
        except (ValueError, IndexError):
            print("错误: --output 参数需要指定输出文件名")
            sys.exit(1)
    
    # 执行转换
    converter = MarkdownToWordConverter()
    try:
        converter.convert(input_file, output_file)
    except Exception as e:
        print(f"错误: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
